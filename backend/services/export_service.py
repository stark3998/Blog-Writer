"""Export Service — Convert MDX/Markdown to multiple output formats.

Supports: Markdown (.md), HTML (.html), PDF (.pdf), DOCX (.docx), MDX (.mdx).
"""

import io
import re
from typing import Literal

import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# WeasyPrint requires system libraries (pango, cairo, glib) that may not be
# available on all platforms (e.g. bare Windows). Import lazily so the rest
# of the export service still works; PDF export will raise a clear error.
WeasyprintHTML = None
try:
    from weasyprint import HTML as WeasyprintHTML
except (ImportError, OSError):
    pass

ExportFormat = Literal["md", "html", "pdf", "docx", "mdx"]

# CSS for HTML and PDF rendering
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
  body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    max-width: 800px;
    margin: 0 auto;
    padding: 2rem;
    line-height: 1.7;
    color: #1a1a1a;
  }}
  h1 {{ font-size: 2rem; margin-top: 2rem; color: #111; }}
  h2 {{ font-size: 1.5rem; margin-top: 1.8rem; color: #222; border-bottom: 1px solid #eee; padding-bottom: 0.3rem; }}
  h3 {{ font-size: 1.25rem; margin-top: 1.5rem; color: #333; }}
  code {{
    background: #f4f4f5;
    padding: 0.15rem 0.4rem;
    border-radius: 3px;
    font-size: 0.9em;
    font-family: 'Fira Code', 'Consolas', monospace;
  }}
  pre {{
    background: #1e1e2e;
    color: #cdd6f4;
    padding: 1rem;
    border-radius: 8px;
    overflow-x: auto;
    line-height: 1.5;
  }}
  pre code {{
    background: none;
    padding: 0;
    color: inherit;
  }}
  blockquote {{
    border-left: 4px solid #6366f1;
    margin: 1rem 0;
    padding: 0.5rem 1rem;
    background: #f8f8ff;
    color: #555;
  }}
  img {{ max-width: 100%; border-radius: 8px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
  th, td {{ border: 1px solid #ddd; padding: 0.5rem 0.75rem; text-align: left; }}
  th {{ background: #f4f4f5; font-weight: 600; }}
  a {{ color: #6366f1; text-decoration: none; }}
  a:hover {{ text-decoration: underline; }}
  .frontmatter {{ background: #f8f9fa; padding: 1rem; border-radius: 8px; margin-bottom: 2rem; border: 1px solid #e9ecef; }}
  .frontmatter p {{ margin: 0.25rem 0; }}
</style>
</head>
<body>
{content}
</body>
</html>"""


def _strip_frontmatter(content: str) -> tuple[dict[str, str], str]:
    """Extract YAML frontmatter and return (metadata, body)."""
    metadata: dict[str, str] = {}
    body = content

    fm_match = re.match(r"^---\s*\n(.*?)\n---\s*\n", content, re.DOTALL)
    if fm_match:
        fm_text = fm_match.group(1)
        body = content[fm_match.end():]

        for line in fm_text.strip().split("\n"):
            kv = line.split(":", 1)
            if len(kv) == 2:
                key = kv[0].strip()
                val = kv[1].strip().strip('"').strip("'")
                metadata[key] = val

    return metadata, body


def _strip_mdx_components(content: str) -> str:
    """Remove MDX-specific JSX components, leaving only standard markdown."""
    # Remove import statements
    content = re.sub(r"^import\s+.*$", "", content, flags=re.MULTILINE)
    # Remove self-closing JSX tags like <Component />
    content = re.sub(r"<[A-Z][a-zA-Z]*\s*[^>]*/\s*>", "", content)
    # Remove opening/closing JSX tags like <Component>...</Component>
    content = re.sub(r"</?[A-Z][a-zA-Z]*[^>]*>", "", content)
    # Clean up excess blank lines
    content = re.sub(r"\n{3,}", "\n\n", content)
    return content.strip()


def _convert_to_html(content: str) -> str:
    """Convert markdown content to styled HTML page."""
    metadata, body = _strip_frontmatter(content)
    body = _strip_mdx_components(body)

    # Convert markdown to HTML
    md = markdown.Markdown(
        extensions=["tables", "fenced_code", "codehilite", "toc", "nl2br"]
    )
    html_body = md.convert(body)

    # Add frontmatter banner if we have metadata
    if metadata:
        fm_html = '<div class="frontmatter">'
        if "title" in metadata:
            fm_html += f"<h1>{metadata['title']}</h1>"
        if "date" in metadata:
            fm_html += f"<p><strong>Date:</strong> {metadata['date']}</p>"
        if "excerpt" in metadata:
            fm_html += f"<p><em>{metadata['excerpt']}</em></p>"
        if "tags" in metadata:
            fm_html += f"<p><strong>Tags:</strong> {metadata['tags']}</p>"
        fm_html += "</div>"
        html_body = fm_html + html_body

    title = metadata.get("title", "Blog Post")
    return HTML_TEMPLATE.format(title=title, content=html_body)


def _convert_to_markdown(content: str) -> str:
    """Convert MDX content to clean Markdown."""
    metadata, body = _strip_frontmatter(content)
    body = _strip_mdx_components(body)

    # Re-add frontmatter as markdown metadata
    parts = []
    if metadata:
        parts.append("---")
        for k, v in metadata.items():
            parts.append(f"{k}: {v}")
        parts.append("---")
        parts.append("")

    parts.append(body)
    return "\n".join(parts)


def _convert_to_pdf(content: str) -> bytes:
    """Convert content to PDF via HTML intermediate."""
    if WeasyprintHTML is None:
        raise RuntimeError(
            "PDF export requires WeasyPrint with system libraries "
            "(pango, cairo, glib). Install them or use the Docker image."
        )
    html_content = _convert_to_html(content)
    pdf_bytes = WeasyprintHTML(string=html_content).write_pdf()
    return pdf_bytes


def _convert_to_docx(content: str) -> bytes:
    """Convert content to DOCX."""
    metadata, body = _strip_frontmatter(content)
    body = _strip_mdx_components(body)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Add title from frontmatter
    title = metadata.get("title", "Blog Post")
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    if metadata.get("date"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Date: {metadata['date']}")
        run.font.size = Pt(10)
        run.font.color.rgb = None  # Use default color

    if metadata.get("excerpt"):
        p = doc.add_paragraph()
        run = p.add_run(metadata["excerpt"])
        run.italic = True

    doc.add_paragraph("")  # Spacer

    # Parse markdown body into document elements
    lines = body.split("\n")
    in_code_block = False
    code_buffer: list[str] = []

    for line in lines:
        # Handle code blocks
        if line.startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s*", "", line)
            doc.add_paragraph(text.strip(), style="List Number")
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line[2:].strip())
            run.italic = True
        elif line.strip():
            # Handle bold and inline code in normal paragraphs
            p = doc.add_paragraph()
            # Simple parsing for **bold** and `code`
            remaining = line
            while remaining:
                bold_match = re.search(r"\*\*(.+?)\*\*", remaining)
                code_match = re.search(r"`(.+?)`", remaining)

                next_match = None
                match_type = None

                if bold_match and code_match:
                    if bold_match.start() < code_match.start():
                        next_match = bold_match
                        match_type = "bold"
                    else:
                        next_match = code_match
                        match_type = "code"
                elif bold_match:
                    next_match = bold_match
                    match_type = "bold"
                elif code_match:
                    next_match = code_match
                    match_type = "code"

                if next_match:
                    # Add text before match
                    before = remaining[: next_match.start()]
                    if before:
                        p.add_run(before)

                    if match_type == "bold":
                        run = p.add_run(next_match.group(1))
                        run.bold = True
                    elif match_type == "code":
                        run = p.add_run(next_match.group(1))
                        run.font.name = "Consolas"
                        run.font.size = Pt(10)

                    remaining = remaining[next_match.end():]
                else:
                    if remaining:
                        p.add_run(remaining)
                    remaining = ""

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_blog(content: str, format: ExportFormat) -> tuple[bytes, str, str]:
    """Export blog content to the specified format.

    Args:
        content: The MDX/Markdown blog content.
        format: Target format ('md', 'html', 'pdf', 'docx', 'mdx').

    Returns:
        Tuple of (file_bytes, filename, content_type).
    """
    metadata, _ = _strip_frontmatter(content)
    slug = metadata.get("title", "blog-post").lower()
    slug = re.sub(r"[^a-z0-9]+", "-", slug).strip("-")[:60]

    if format == "mdx":
        return (
            content.encode("utf-8"),
            f"{slug}.mdx",
            "text/mdx",
        )
    elif format == "md":
        md_content = _convert_to_markdown(content)
        return (
            md_content.encode("utf-8"),
            f"{slug}.md",
            "text/markdown",
        )
    elif format == "html":
        html_content = _convert_to_html(content)
        return (
            html_content.encode("utf-8"),
            f"{slug}.html",
            "text/html",
        )
    elif format == "pdf":
        pdf_bytes = _convert_to_pdf(content)
        return (
            pdf_bytes,
            f"{slug}.pdf",
            "application/pdf",
        )
    elif format == "docx":
        docx_bytes = _convert_to_docx(content)
        return (
            docx_bytes,
            f"{slug}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        raise ValueError(f"Unsupported export format: {format}")
